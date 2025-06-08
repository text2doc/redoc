"""DOCX converter implementation for Redoc."""

import json
import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from . import register_converter, BaseConverter
from ..exceptions import ConversionError


@register_converter('docx')
class DocxConverter(BaseConverter):
    """Converter for Microsoft Word (DOCX) documents."""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'html', 'txt', 'md', 'odt', 'rtf']
    
    def convert(
        self,
        source: Union[str, Path, Dict],
        output_file: Optional[Union[str, Path]] = None,
        **kwargs
    ) -> Any:
        """Convert to/from DOCX format.
        
        Args:
            source: Source file path or template dictionary
            output_file: Optional output file path
            **kwargs: Additional conversion options
                - from_format: Source format if not inferrable from source
                - to_format: Target format if not inferrable from output_file
                - style: Document style for conversion
        """
        try:
            # Handle template input
            if isinstance(source, dict):
                return self._convert_from_template(source, output_file, **kwargs)
                
            source_path = Path(source)
            
            # If output_file is not specified, determine from source
            if output_file is None:
                if 'to_format' not in kwargs:
                    raise ValueError("Either output_file or to_format must be specified")
                to_format = kwargs['to_format']
                output_file = source_path.with_suffix(f'.{to_format}')
            
            output_path = Path(output_file)
            from_format = kwargs.get('from_format', 'docx' if source_path.suffix.lower() == '.docx' else None)
            to_format = kwargs.get('to_format', output_path.suffix[1:].lower())
            
            if from_format == 'docx':
                return self._convert_from_docx(source_path, output_path, to_format, **kwargs)
            else:
                return self._convert_to_docx(source_path, output_path, from_format, **kwargs)
                
        except Exception as e:
            raise ConversionError(f"DOCX conversion failed: {str(e)}") from e
    
    def _convert_from_docx(
        self,
        source_path: Path,
        output_path: Path,
        to_format: str,
        **kwargs
    ) -> Path:
        """Convert from DOCX to another format."""
        try:
            from docx import Document
            
            # Read the DOCX file
            doc = Document(source_path)
            
            if to_format == 'pdf':
                # Convert DOCX to PDF using unoconv or libreoffice
                return self._docx_to_pdf(source_path, output_path, **kwargs)
                
            elif to_format == 'html':
                # Convert DOCX to HTML
                html = self._docx_to_html(doc, **kwargs)
                output_path.parent.mkdir(parents=True, exist_ok=True)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html)
                return output_path
                
            elif to_format == 'txt':
                # Convert DOCX to plain text
                text = self._docx_to_text(doc)
                output_path.parent.mkdir(parents=True, exist_ok=True)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                return output_path
                
            elif to_format == 'md':
                # Convert DOCX to Markdown
                markdown = self._docx_to_markdown(doc, **kwargs)
                output_path.parent.mkdir(parents=True, exist_ok=True)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(markdown)
                return output_path
                
            elif to_format in ['odt', 'rtf']:
                # Use libreoffice for format conversion
                return self._convert_with_libreoffice(source_path, output_path, to_format, **kwargs)
                
            else:
                raise ConversionError(f"Conversion from DOCX to {to_format} not supported")
            
        except ImportError as e:
            raise ConversionError(f"Required dependency not found: {str(e)}") from e
        except Exception as e:
            raise ConversionError(f"Failed to convert DOCX to {to_format}: {str(e)}") from e
    
    def _convert_to_docx(
        self,
        source_path: Path,
        output_path: Path,
        from_format: str,
        **kwargs
    ) -> Path:
        """Convert from another format to DOCX."""
        try:
            if from_format == 'html':
                # Convert HTML to DOCX
                return self._html_to_docx(source_path, output_path, **kwargs)
                
            elif from_format == 'md':
                # Convert Markdown to DOCX
                return self._markdown_to_docx(source_path, output_path, **kwargs)
                
            elif from_format in ['odt', 'rtf', 'txt']:
                # Use libreoffice for format conversion
                return self._convert_with_libreoffice(source_path, output_path, 'docx', **kwargs)
                
            else:
                raise ConversionError(f"Conversion from {from_format} to DOCX not implemented")
                
        except Exception as e:
            raise ConversionError(f"Failed to convert {from_format} to DOCX: {str(e)}") from e
    
    def _docx_to_pdf(
        self,
        source_path: Path,
        output_path: Path,
        **kwargs
    ) -> Path:
        """Convert DOCX to PDF using unoconv or libreoffice."""
        try:
            # First try unoconv
            try:
                import unoconv
                unoconv.convert(str(source_path), output=str(output_path), format='pdf')
                return output_path
            except (ImportError, Exception):
                pass
            
            # Fall back to libreoffice
            try:
                import subprocess
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', str(output_path.parent),
                    str(source_path)
                ]
                subprocess.run(cmd, check=True, capture_output=True, text=True)
                
                # Rename the output file if needed
                pdf_path = source_path.with_suffix('.pdf')
                if pdf_path != output_path:
                    pdf_path.replace(output_path)
                
                return output_path
            except Exception as e:
                raise ConversionError(
                    "Failed to convert DOCX to PDF. "
                    "Install unoconv (pip install unoconv) or libreoffice."
                ) from e
                
        except Exception as e:
            raise ConversionError(f"Failed to convert DOCX to PDF: {str(e)}") from e
    
    def _docx_to_html(
        self,
        doc: 'Document',
        **kwargs
    ) -> str:
        """Convert a python-docx Document to HTML."""
        from bs4 import BeautifulSoup
        
        # Create a new BeautifulSoup document
        soup = BeautifulSoup('', 'html.parser')
        html = soup.new_tag('html')
        soup.append(html)
        
        # Add head and body
        head = soup.new_tag('head')
        html.append(head)
        
        title = soup.new_tag('title')
        title.string = 'Converted Document'
        head.append(title)
        
        # Add basic CSS
        style = soup.new_tag('style')
        style.string = """
            body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }
            h1, h2, h3 { color: #2c3e50; }
            p { margin: 10px 0; }
            table { border-collapse: collapse; width: 100%; margin: 20px 0; }
            th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            th { background-color: #f2f2f2; }
            img { max-width: 100%; height: auto; }
        """
        head.append(style)
        
        body = soup.new_tag('body')
        html.append(body)
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.lstrip('Heading '))
                h_tag = soup.new_tag(f'h{min(6, level)}')
                h_tag.string = para.text
                body.append(h_tag)
            else:
                p = soup.new_tag('p')
                p.string = para.text
                body.append(p)
        
        # Process tables
        for table in doc.tables:
            table_html = soup.new_tag('table')
            
            # Add header row if exists
            if len(table.rows) > 0 and any(cell.text.strip() for cell in table.rows[0].cells):
                thead = soup.new_tag('thead')
                tr = soup.new_tag('tr')
                for cell in table.rows[0].cells:
                    th = soup.new_tag('th')
                    th.string = cell.text
                    tr.append(th)
                thead.append(tr)
                table_html.append(thead)
                    
                # Add body rows
                tbody = soup.new_tag('tbody')
                for row in table.rows[1:]:
                        tr = soup.new_tag('tr')
                        for cell in row.cells:
                            td = soup.new_tag('td')
                            td.string = cell.text
                            tr.append(td)
                        tbody.append(tr)
                table_html.append(tbody)
            
            body.append(table_html)
        
        return soup.prettify()
    
    def _docx_to_text(
        self,
        doc: 'Document',
        **kwargs
    ) -> str:
        """Convert a python-docx Document to plain text."""
        lines = []
        
        for para in doc.paragraphs:
            lines.append(para.text)
        
        # Add tables as tab-separated values
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                lines.append('\t'.join(row_text))
            lines.append('')  # Add empty line after each table
        
        return '\n'.join(lines)
    
    def _docx_to_markdown(
        self,
        doc: 'Document',
        **kwargs
    ) -> str:
        """Convert a python-docx Document to Markdown."""
        from bs4 import BeautifulSoup
        
        # First convert to HTML
        html = self._docx_to_html(doc, **kwargs)
        
        # Then convert HTML to Markdown
        try:
            import markdownify
            return markdownify.markdownify(html, heading_style='ATX')
        except ImportError:
            raise ConversionError(
                "markdownify package is required for DOCX to Markdown conversion. "
                "Install with: pip install markdownify"
            )
    
    def _html_to_docx(
        self,
        source_path: Path,
        output_path: Path,
        **kwargs
    ) -> Path:
        """Convert HTML to DOCX."""
        try:
            from htmldocx import HtmlToDocx
            
            # Read HTML content
            with open(source_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Convert HTML to DOCX
            docx_converter = HtmlToDocx()
            docx = docx_converter.parse_html_string(html_content)
            
            # Save the document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            docx.save(output_path)
            
            return output_path
            
        except ImportError as e:
            raise ConversionError(
                "htmldocx package is required for HTML to DOCX conversion. "
                "Install with: pip install htmldocx"
            ) from e
        except Exception as e:
            raise ConversionError(f"Failed to convert HTML to DOCX: {str(e)}") from e
    
    def _markdown_to_docx(
        self,
        source_path: Path,
        output_path: Path,
        **kwargs
    ) -> Path:
        """Convert Markdown to DOCX."""
        try:
            import markdown
            from bs4 import BeautifulSoup
            
            # Read Markdown content
            with open(source_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Convert Markdown to HTML
            html = markdown.markdown(md_content)
            
            # Save HTML to a temporary file
            with tempfile.NamedTemporaryFile(suffix='.html', delete=False, mode='w', encoding='utf-8') as f:
                f.write(html)
                temp_html_path = f.name
            
            try:
                # Convert HTML to DOCX
                self._html_to_docx(Path(temp_html_path), output_path, **kwargs)
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_html_path)
                except:
                    pass
            
            return output_path
            
        except ImportError as e:
            raise ConversionError(
                "markdown package is required for Markdown to DOCX conversion. "
                "Install with: pip install markdown"
            ) from e
        except Exception as e:
            raise ConversionError(f"Failed to convert Markdown to DOCX: {str(e)}") from e
    
    def _convert_with_libreoffice(
        self,
        source_path: Path,
        output_path: Path,
        target_format: str,
        **kwargs
    ) -> Path:
        """Convert between formats using libreoffice."""
        try:
            import subprocess
            
            # Map target format to libreoffice format
            format_map = {
                'pdf': 'pdf',
                'docx': 'docx',
                'odt': 'odt',
                'rtf': 'rtf',
                'txt': 'txt',
                'html': 'html',
            }
            
            if target_format not in format_map:
                raise ConversionError(f"Unsupported target format for libreoffice conversion: {target_format}")
            
            # Create output directory if it doesn't exist
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Run libreoffice for conversion
            cmd = [
                'libreoffice', '--headless', '--convert-to', format_map[target_format],
                '--outdir', str(output_path.parent),
                str(source_path)
            ]
            
            result = subprocess.run(cmd, check=True, capture_output=True, text=True)
            
            # Rename the output file if needed
            output_ext = f'.{target_format}'
            default_output = source_path.with_suffix(output_ext)
            
            if default_output != output_path:
                default_output.replace(output_path)
            
            return output_path
            
        except subprocess.CalledProcessError as e:
            raise ConversionError(
                f"libreoffice conversion failed with return code {e.returncode}: {e.stderr}"
            ) from e
        except Exception as e:
            raise ConversionError(f"Failed to convert with libreoffice: {str(e)}") from e
    
    def _convert_from_template(
        self,
        template: Dict[str, Any],
        output_file: Optional[Union[str, Path]] = None,
        **kwargs
    ) -> Path:
        """Convert a template to DOCX."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create a new document
            doc = Document()
            
            # Apply template data
            if 'title' in template:
                doc.add_heading(template['title'], level=1)
            
            if 'content' in template:
                if isinstance(template['content'], str):
                    doc.add_paragraph(template['content'])
                elif isinstance(template['content'], list):
                    for item in template['content']:
                        if isinstance(item, dict):
                            if item.get('type') == 'heading':
                                doc.add_heading(item.get('text', ''), level=item.get('level', 1))
                            elif item.get('type') == 'paragraph':
                                p = doc.add_paragraph()
                                if 'style' in item:
                                    p.style = item['style']
                                p.add_run(item.get('text', ''))
                            elif item.get('type') == 'table':
                                self._add_table(doc, item)
                        else:
                            doc.add_paragraph(str(item))
            
            # Save the document
            if output_file is None:
                output_file = 'output.docx'
            
            output_path = Path(output_file)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            
            return output_path
            
        except ImportError as e:
            raise ConversionError(
                "python-docx package is required for template conversion. "
                "Install with: pip install python-docx"
            ) from e
        except Exception as e:
            raise ConversionError(f"Template conversion failed: {str(e)}") from e
    
    def _add_table(self, doc: 'Document', table_data: Dict[str, Any]) -> None:
        """Add a table to a document from table data."""
        if 'data' not in table_data:
            return
        
        # Create table
        rows = table_data['data']
        if not rows:
            return
            
        cols = max(len(row) for row in rows) if rows else 0
        table = doc.add_table(rows=len(rows), cols=cols)
        
        # Apply table style if specified
        if 'style' in table_data:
            table.style = table_data['style']
        
        # Add data to table
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                table.cell(i, j).text = str(cell)
    
    def get_supported_formats(self) -> list:
        """Get a list of formats this converter can handle."""
        return self.supported_formats
